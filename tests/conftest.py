"""Shared fixtures for tests."""

import asyncio
import tempfile
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock

import pytest
import pytest_asyncio


@pytest.fixture
def tmp_dir():
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


@pytest.fixture
def sample_pdf(tmp_dir):
    """Create a minimal valid PDF for testing."""
    pdf_path = tmp_dir / "test.pdf"
    # Minimal PDF structure
    pdf_content = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<<>>>>endobj
xref
0 4
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
trailer<</Size 4/Root 1 0 R>>
startxref
206
%%EOF"""
    pdf_path.write_bytes(pdf_content)
    return pdf_path


@pytest.fixture
def sample_text_file(tmp_dir):
    path = tmp_dir / "test.txt"
    path.write_text("This is a test document about medical diagnosis and blood test results.")
    return path


@pytest.fixture
def sample_docx(tmp_dir):
    """Create a minimal DOCX for testing (if python-docx available)."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is a test paragraph about business invoices and payments.")
        path = tmp_dir / "test.docx"
        doc.save(str(path))
        return path
    except ImportError:
        return None


@pytest_asyncio.fixture
async def db(tmp_dir):
    from app.storage.db import Database
    database = Database(tmp_dir / "test.db")
    await database.connect()
    yield database
    await database.close()


@pytest.fixture
def file_storage(tmp_dir):
    from app.storage.files import FileStorage
    return FileStorage(base_path=tmp_dir / "files", allowed_extensions=[".pdf", ".txt", ".docx", ".jpg", ".png"])


@pytest.fixture
def mock_llm_router():
    from app.llm.router import LLMResponse
    router = MagicMock()
    router.classify = AsyncMock(return_value=LLMResponse(
        text='{"category":"health","confidence":0.9,"tags":["medical"],"summary":"Test doc","document_type":"lab_result"}',
        model="test-model", role="classification",
    ))
    router.get_stats = MagicMock(return_value={"total_calls": 0, "total_cost_usd": 0})
    return router
