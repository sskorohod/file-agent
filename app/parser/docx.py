"""DOCX parser — extract text preserving structure."""

from __future__ import annotations

import logging
from pathlib import Path

from .base import ParseResult

logger = logging.getLogger(__name__)


class DocxParser:
    """Parse DOCX files using python-docx."""

    supported_extensions = [".docx"]

    def can_handle(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.supported_extensions

    async def parse(self, file_path: Path) -> ParseResult:
        from docx import Document

        doc = Document(str(file_path))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading structure
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    parts.append(f"{'#' * int(level)} {text}" if level.isdigit() else text)
                else:
                    parts.append(text)

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))

        text = "\n\n".join(parts)

        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }
        # Core properties
        try:
            cp = doc.core_properties
            if cp.author:
                metadata["author"] = cp.author
            if cp.title:
                metadata["title"] = cp.title
            if cp.created:
                metadata["created"] = cp.created.isoformat()
        except Exception:
            pass

        return ParseResult(
            text=text,
            metadata=metadata,
            pages=0,  # DOCX doesn't have pages natively
            language=self._detect_language(text),
            parser_used="python-docx",
        )

    def _detect_language(self, text: str) -> str:
        if not text:
            return "unknown"
        sample = text[:2000]
        cyrillic = sum(1 for c in sample if "\u0400" <= c <= "\u04ff")
        latin = sum(1 for c in sample if "a" <= c.lower() <= "z")
        total = cyrillic + latin
        if total == 0:
            return "unknown"
        return "ru" if cyrillic / total > 0.5 else "en"


class TextParser:
    """Parse plain text files (TXT, CSV, etc.)."""

    supported_extensions = [".txt", ".csv", ".tsv", ".md", ".json", ".xml"]

    def can_handle(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.supported_extensions

    async def parse(self, file_path: Path) -> ParseResult:
        try:
            text = file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = file_path.read_text(encoding="latin-1")

        return ParseResult(
            text=text,
            metadata={"encoding": "utf-8", "lines": text.count("\n") + 1},
            pages=0,
            language="",
            parser_used="text",
        )
